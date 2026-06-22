from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

NAVY = RGBColor(0x14, 0x23, 0x3D)
RED = RGBColor(0xD7, 0x26, 0x26)
MUTED = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_para(text="", bold=False, size=11, color=None, align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color or NAVY
    return p


def add_h2(text):
    p = add_para(text, bold=True, size=14, color=RED, space_before=14, space_after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "D72626")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h3(text, color=None):
    return add_para(text, bold=True, size=11, color=color or NAVY, space_before=10, space_after=2)


def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    r.italic = True
    return p


def add_body(text):
    return add_para(text, size=10.5, color=NAVY, space_before=2, space_after=6)


def make_table(headers, rows, col_widths=None, header_bg="142340"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after = Pt(3)
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        bg = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = NAVY
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def make_done_table(headers, rows, col_widths=None):
    return make_table(headers, rows, col_widths, header_bg="065F46")


# ─── Title ──────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(2)
r = title.add_run("NxtSft")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
r = sub.add_run("PRD Gap Analysis — What Remains to Build")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RED

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(6)
r = meta.add_run(f"Updated: {datetime.date.today().strftime('%B %d, %Y')}   ·   Branch: main   ·   Commit: 7e43bc4   ·   Status: v1.0 demo phase")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

hr = doc.add_paragraph()
hr.paragraph_format.space_after = Pt(12)
pPr = hr._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "D72626")
pBdr.append(bottom); pPr.append(pBdr)

add_h3("Executive Summary")
add_body(
    "All portal routes and public pages exist. Since the last revision (June 17), the Maps integration, "
    "the Supervisor Visit Calendar (now live data + geographic map), the Role & Permission Matrix, the "
    "property gallery lightbox, notification-preference persistence, and the DB-backed agents directory "
    "have all been completed. The remaining work concentrates in three areas: (1) payment + communications "
    "integrations (Razorpay, SMS/WhatsApp/Email/VoIP); (2) three advanced admin tabs still on static data "
    "(Marketing, AI Model Control, Content CMS, Billing); and (3) a handful of smaller UI/polish items."
)

# ─── Completed since last revision ─────────────────────────────────────────
add_h2("Completed Since Last Revision  (June 17 → 19)")
make_done_table(
    headers=["Area", "Feature", "How It Was Closed"],
    rows=[
        ["Maps", "Property location pin + Supervisor visit map",
         "Mapbox via react-map-gl. PropertyMap (single pin on /properties/[slug]) and VisitsMap (geographic, colour-coded per sales rep). Shared lib/map.ts with city-centroid fallback for listings without coordinates."],
        ["Supervisor Portal", "Visit Calendar — live data",
         "New siteVisits.mapData staff query (batch-joins property coords + rep). Tab rebuilt from fixtures to live data: stat row + map + visit list."],
        ["Super Admin", "Role & Permission Matrix",
         "4-level access grid (none/read/write/admin) over 12 features x 6 roles, click-to-cycle cells, plus a 'simulate as role' tester. Persisted via superAdmin.getPermissionMatrix / updatePermissionMatrix."],
        ["Property Detail", "Gallery lightbox",
         "GalleryLightbox component — full-screen image view with navigation."],
        ["User Profile", "Notification preferences persistence",
         "Wired to users.notificationPrefs query + updateNotificationPrefs mutation — toggles now persist per user."],
        ["Public Site", "Agents directory DB-backed",
         "/agents now fetches via trpcClient.users.getAgents (live PostgreSQL) instead of a static fixture."],
        ["Super Admin", "Security Console",
         "Confirmed already wired to tRPC (failed logins, security log, IP allow/blocklist, password/2FA policy). Previously mislabelled as static."],
    ],
    col_widths=[3.0, 4.2, 6.6],
)

# ─── P0 ─────────────────────────────────────────────────────────────────────
add_h2("P0 — Blocks Launch")
make_table(
    headers=["Item", "Current State", "What Is Needed"],
    rows=[
        ["Razorpay — real checkout",
         "Backend createOrder / verifyPayment exist. Frontend passes demo payment IDs.",
         "Load Razorpay checkout.js, open the payment modal, pass the real order ID. Applies to credit packs and subscription purchases."],
        ["SMS OTP for 2FA",
         "2FA toggle flips the DB flag but sends no OTP — no phone verification.",
         "SMS provider (Twilio / MSG91): OTP send + verify endpoint wired into the 2FA enable flow."],
    ],
    col_widths=[3.3, 5.0, 5.5],
)

# ─── P1 ─────────────────────────────────────────────────────────────────────
add_h2("P1 — Core Feature Completeness")
add_body("Screens exist and render; the listed interaction or data wiring is missing.")
make_table(
    headers=["Feature", "Portal / Route", "Gap"],
    rows=[
        ["CRM Pipeline — drag & drop", "Admin",
         "Kanban columns exist; stage moves use a Select dropdown. Add card drag-and-drop (no dnd library installed yet)."],
        ["Marketing — campaign builder", "Admin",
         "Static campaign table with pause/resume. Build create-campaign form (type, audience, subject/body, schedule, send) + live campaigns data."],
        ["Gallery drag-to-reorder", "/list",
         "Multi-image upload works; add a drag handle to reorder images before submit."],
        ["Recent Activity — live data", "/profile",
         "Activity timeline shows 5 hardcoded entries. Wire to a real audit/activity log."],
        ["Agent profile page", "/agents/[slug]",
         "Directory is DB-backed, but the individual profile still uses the AGENTS fixture + static properties. Wire to live data."],
    ],
    col_widths=[3.6, 2.6, 7.6],
    header_bg="1A3A5C",
)

# ─── P2 ─────────────────────────────────────────────────────────────────────
add_h2("P2 — Advanced Admin Features (static → build out)")
make_table(
    headers=["Tab", "Current State", "Gap"],
    rows=[
        ["AI Model Control (Super Admin)",
         "Static status table; deploy button is a toast.",
         "Feature-weight sliders (location/price/BHK), match-score calibration, drift trend chart, rollback + live model_versions data."],
        ["Content CMS (Super Admin)",
         "Static page list; '+ New Page' is a toast.",
         "Rich-text editor (TipTap / Quill), media library, blog post CRUD, schedule/publish workflow + live cms_pages data."],
        ["Billing & Revenue (Super Admin)",
         "Static data.",
         "Invoice PDF generation + download, GST-compliant tax report + live invoices data."],
    ],
    col_widths=[3.6, 4.4, 6.0],
    header_bg="1A3A5C",
)

# ─── P3 ─────────────────────────────────────────────────────────────────────
add_h2("P3 — Integration Infrastructure (vendor onboarding)")
make_table(
    headers=["Integration", "Use Cases"],
    rows=[
        ["WhatsApp Business API", "Bulk blast campaigns (Marketing) + one-tap lead contact shortcut."],
        ["Email service (Resend / SendGrid)", "Transactional emails (listing approved, payment receipt, subscription renewal) + marketing campaign sends."],
        ["VoIP / PSTN (Exotel / Twilio)", "Sales Click-to-Call: outbound call trigger, live timer, post-call note auto-sync. UI is complete; 'Dial' is currently a toast."],
    ],
    col_widths=[4.0, 9.8],
)

# ─── P4 ─────────────────────────────────────────────────────────────────────
add_h2("P4 — Polish / Verify")
make_table(
    headers=["Item", "Note"],
    rows=[
        ["Property search filters DB-backed", "Spot-check that price/BHK/city/type filters hit PostgreSQL."],
        ["Legal pages copy", "Confirm /terms, /privacy, /cookie-policy, /fraud-advisory carry final legal text, not placeholder."],
        ["Home KPI count-up animation", "Marked done previously — verify it fires on scroll-in."],
    ],
    col_widths=[4.6, 9.2],
)

# ─── Decision ───────────────────────────────────────────────────────────────
add_h2("Not in PRD — Product Decision Required")
make_table(
    headers=["Route", "Description", "Action Needed"],
    rows=[
        ["/owners/[slug]", "Owner profile page", "Define PRD spec or remove from nav/sitemap."],
        ["/refer", "Referral / invite page (a public route)", "Define referral mechanics (reward, tracking, T&C) or remove."],
    ],
    col_widths=[3.0, 4.5, 6.3],
    header_bg="6B2424",
)

# ─── Summary ────────────────────────────────────────────────────────────────
add_h2("Priority Summary")
make_table(
    headers=["Priority", "Work Area", "Estimated Effort"],
    rows=[
        ["P0 — Blocks launch", "Razorpay checkout, SMS OTP for 2FA", "2–4 days"],
        ["P1 — Core completeness", "CRM drag-drop, campaign builder, gallery reorder, recent-activity + agent-profile live data", "1–1.5 weeks"],
        ["P2 — Advanced admin", "AI Model Control, Content CMS, Billing invoices", "2–3 weeks"],
        ["P3 — Integrations", "WhatsApp, Email, VoIP", "Depends on vendor onboarding"],
        ["P4 — Polish", "Search filters, legal copy, animation checks", "1–2 days"],
    ],
    col_widths=[3.6, 7.7, 4.5],
)

add_note("Source of truth: docs/nxtsft_prd.pdf. This document reflects commit 7e43bc4 on main, June 19, 2026.")

out = r"d:\LookAround\8. Nxtsft\docs\NxtSft_PRD_Gap_Analysis.docx"
doc.save(out)
print(f"Saved: {out}")
